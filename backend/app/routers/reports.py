from fastapi import APIRouter, HTTPException
from pydantic import BaseModel
from docx import Document
from fpdf import FPDF
import io
from fastapi.responses import FileResponse

router = APIRouter(prefix="/reports", tags=["reports"])

class ReportRequest(BaseModel):
    title: str
    content: str
    format: str  # "pdf" or "docx"

@router.post("/generate")
async def generate_report(report: ReportRequest):
    try:
        if report.format.lower() == "pdf":
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            pdf.multi_cell(0, 10, report.content)
            file_path = f"{report.title}.pdf"
            pdf.output(file_path)
        elif report.format.lower() == "docx":
            doc = Document()
            doc.add_heading(report.title, 0)
            doc.add_paragraph(report.content)
            file_path = f"{report.title}.docx"
            doc.save(file_path)
        else:
            raise HTTPException(status_code=400, detail="Invalid format. Use 'pdf' or 'docx'.")
        return FileResponse(path=file_path, filename=file_path)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
